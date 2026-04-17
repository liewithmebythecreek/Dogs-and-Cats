"""
generate_doc.py — generates NeuralWeather_Project_Presentation.docx
Run: python generate_doc.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml   import OxmlElement
import copy

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY    = RGBColor(0x0B, 0x0F, 0x1A)   # dark background
ORANGE  = RGBColor(0xF9, 0x73, 0x16)   # ECMWF accent
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
SILVER  = RGBColor(0x94, 0xA3, 0xB8)
DARK    = RGBColor(0x1E, 0x25, 0x40)   # slide body dark


def set_cell_bg(cell, hex_color: str):
    """Fill a table cell with a solid background colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def add_page_break(doc):
    doc.add_page_break()


def slide_title(doc, title: str, subtitle: str = ""):
    """Full-width dark title block with optional subtitle."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    cell.width = Inches(9)
    set_cell_bg(cell, '0B0F1A')

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after  = Pt(8)
    run = p.add_run(title)
    run.bold      = True
    run.font.size = Pt(30)
    run.font.color.rgb = ORANGE

    if subtitle:
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(28)
        r2 = p2.add_run(subtitle)
        r2.font.size = Pt(14)
        r2.font.color.rgb = SILVER

    doc.add_paragraph()   # breathing space


def section_heading(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(14)
    run.font.color.rgb = ORANGE


def body_bullet(doc, text: str, level: int = 0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent   = Inches(0.3 * (level + 1))
    p.paragraph_format.space_after   = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(11)


def two_col_table(doc, headers, rows, col_widths=None):
    """Generic table with a dark header row."""
    n_cols = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=n_cols)
    t.style = 'Table Grid'

    # Header
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, '1E2540')
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = ORANGE
        r.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for ri, row in enumerate(rows):
        tr = t.rows[ri + 1]
        bg = 'F8FAFC' if ri % 2 == 0 else 'EFF6FF'
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)

    doc.add_paragraph()


def add_code_block(doc, code: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x0F, 0x77, 0x2A)


# ═════════════════════════════════════════════════════════════════════════════
#  BUILD DOCUMENT
# ═════════════════════════════════════════════════════════════════════════════

doc = Document()

# Page margins
sec = doc.sections[0]
sec.page_width   = Inches(10)
sec.page_height  = Inches(7.5)
sec.left_margin  = sec.right_margin  = Inches(0.6)
sec.top_margin   = sec.bottom_margin = Inches(0.4)

# ─────────────────────────────────────────────────────────────────────────────
# SLIDE 1 — Title
# ─────────────────────────────────────────────────────────────────────────────
slide_title(doc,
    "NeuralWeather",
    "A GraphCast-Inspired Spatio-Temporal Graph Neural Network\nfor 48-Hour Multi-City Weather Forecasting"
)

two_col_table(doc,
    ['Field', 'Detail'],
    [
        ['Author',        'Student — 4th Year B.Tech / BS (Data Science / CSE)'],
        ['Stack',         'PyTorch · PyTorch Geometric · FastAPI · React · Vite'],
        ['Model',         'Graph-LSTM (STGNN) — 7 cities, 48-h autoregressive horizon'],
        ['Inspiration',   'ECMWF GraphCast Charts (charts.ecmwf.int)'],
        ['Deployment',    'Local i5 machine · GitHub Actions (self-hosted runner)'],
        ['Data Source',   'Open-Meteo API (ERA5-equivalent, free & open-source)'],
    ]
)

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# SLIDE 2 — Problem Statement
# ─────────────────────────────────────────────────────────────────────────────
slide_title(doc, "Problem Statement", "Why build this?")

section_heading(doc, "The Gap in Accessible AI Weather Forecasting")
for b in [
    "State-of-the-art ML weather models (GraphCast, Pangu, FourCastNet) require tens of TPUs to train and run.",
    "No open-source, runnable, end-to-end GNN weather stack exists for students or researchers on consumer hardware.",
    "Traditional LSTM forecasters treat each city independently — ignoring spatial correlations (a cold front moves).",
    "Model drift is silent: a model trained in December degrades in summer without anyone noticing.",
]:
    body_bullet(doc, b)

section_heading(doc, "Our Solution")
for b in [
    "A Graph-LSTM that represents 7 Punjab-region cities as graph nodes connected by geographic proximity.",
    "A dark, ECMWF-style interactive dashboard for visualising 48-hour animated forecasts.",
    "An MLOps pipeline for automated drift detection and fine-tuning — runs on a local i5 CPU.",
]:
    body_bullet(doc, b)

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# SLIDE 3 — System Architecture
# ─────────────────────────────────────────────────────────────────────────────
slide_title(doc, "System Architecture", "Full-stack ML platform")

section_heading(doc, "High-Level Data Flow")
add_code_block(doc,
"""Open-Meteo API (live hourly data)
        │
        ▼
  FastAPI Backend  ──────────────────────────────────────────────────────────────
  ├─ /current   ──► fetch_all_live_data()   ──► 7-city DataFrame dict
  ├─ /forecast  ──► ml_service.predict()   ──► STGNN inference ──► 48×7×6 preds
  └─ /health    ──► model load status
        │
        ▼
  React Frontend (Vite + shadcn)
  ├─ ControlPanel sidebar  (city / variable selector)
  ├─ ForecastChart         (animated AreaChart with ReferenceLine needle)
  └─ PlaybackBar           (⏮ ◀ ▶ ⏩ ⏭ scrubber, 0.5×–4× speed, loop)
        │
        ▼ (background)
  Forecast Logger ──────────────────────────────────────────────────────────────
  └─ forecast_log.jsonl  (predictions stored for drift detection)"""
)

two_col_table(doc,
    ['Layer', 'Technology', 'Purpose'],
    [
        ['Data Ingestion',  'Open-Meteo + httpx async',     'Live & historical ERA5-equivalent weather'],
        ['ML Model',        'PyTorch + PyTorch Geometric',  'Graph-LSTM autoregressive 48-h forecast'],
        ['API Server',      'FastAPI + uvicorn',            'REST endpoints · CORS · model singleton'],
        ['Frontend',        'React 18 + Vite + shadcn/ui',  'ECMWF-style dark animated dashboard'],
        ['MLOps',           'GitHub Actions + DVC',         'Drift detection + fine-tuning pipeline'],
        ['Containerisation','Docker + docker-compose',      'One-command local dev environment'],
    ]
)

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# SLIDE 4 — The Graph-LSTM Model
# ─────────────────────────────────────────────────────────────────────────────
slide_title(doc, "The Graph-LSTM Model (STGNN)", "Encoder · Processor · Decoder")

section_heading(doc, "City Graph — 7 Nodes, Proximity Edges (≤ 200 km)")
two_col_table(doc,
    ['City', 'State', 'Elevation (m)', 'Connected to'],
    [
        ['Ropar',      'Punjab',          '260',  'Chandigarh, Ludhiana, Patiala'],
        ['Chandigarh', 'UT',              '321',  'Ropar, Ludhiana, Ambala'],
        ['Ludhiana',   'Punjab',          '242',  'Ropar, Chandigarh, Jalandhar, Patiala'],
        ['Patiala',    'Punjab',          '250',  'Ropar, Ludhiana, Ambala'],
        ['Jalandhar',  'Punjab',          '228',  'Ludhiana'],
        ['Ambala',     'Haryana',         '264',  'Chandigarh, Patiala'],
        ['Shimla',     'Himachal Pradesh','2276', 'Ropar (via threshold)'],
    ]
)

section_heading(doc, "Forward Pass (per inference call)")
for b in [
    "Warm-up: 48 hours of history → Encoder (Linear+ReLU) → 2× DenseGCNConv (spatial) → LSTMCell (temporal)",
    "Hidden state (hx, cx) carries the 'memory' of all 48 warm-up steps.",
    "Autoregressive decode: for each of 48 future hours, run GCN+LSTM → Tanh-bounded residual delta",
    "pred_dyn = (prev_dyn + residual × 0.05).clamp(0, 1)   ← prevents flat-line accumulation",
    "Output: [1, 48, 7, 6] tensor → inverse MinMax-scaled → JSON forecast per city",
]:
    body_bullet(doc, b)

section_heading(doc, "Features per Node per Timestep (11 total)")
body_bullet(doc, "Dynamic (6): temperature_2m · relative_humidity_2m · wind_speed_10m · surface_pressure · precipitation · weather_code", 0)
body_bullet(doc, "Static  (5): sin_hour · cos_hour · elevation · latitude · longitude", 0)

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# SLIDE 5 — Training Pipeline
# ─────────────────────────────────────────────────────────────────────────────
slide_title(doc, "Training Pipeline", "Cold start · Fine-tuning · Data Sync")

two_col_table(doc,
    ['Mode', 'Command', 'LR', 'Patience', 'Data Window', 'Epochs'],
    [
        ['Cold Start',  'python -m backend.training.train_model',                  '1e-3', '5', '365 days', '50 max'],
        ['Fine-tune',   'python -m backend.training.train_model --finetune',       '1e-5', '3', '30 days',  '10 max'],
        ['Freeze GCN',  '... --finetune --freeze-gcn',                             '1e-5', '3', '30 days',  '10 max'],
        ['Custom',      '... --finetune --epochs 5',                               '1e-5', '3', '30 days',  '5'],
    ]
)

section_heading(doc, "Training-Serving Skew Prevention")
body_bullet(doc, "backend/shared_config.py is the single source of truth for DYNAMIC_FEATURES, CITIES, TIME_STEPS, MODEL_PATH, etc.")
body_bullet(doc, "Training (preprocess.py), inference (ml_service.py), live API (weather_api.py), and drift monitor all import from shared_config.")
body_bullet(doc, "Fine-tuning REUSES the existing scaler — it is never refitted on new data (which would shift the MinMax range).")

section_heading(doc, "Loss Function — Time-Weighted MSE")
add_code_block(doc, "weight_t = exp(−t × 0.05)   →   near-future steps penalised more heavily than far-future")

section_heading(doc, "Optimisations for CPU")
for b in [
    "torch.set_num_threads(os.cpu_count()) — uses all available cores",
    "torch.no_grad() during inference — no computation graph built",
    "torch.nn.utils.clip_grad_norm_(max_norm=1.0) — prevents exploding gradients in fine-tuning",
    "MinMaxScaler fitted per city — each node normalised independently",
]:
    body_bullet(doc, b)

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# SLIDE 6 — Frontend / UI
# ─────────────────────────────────────────────────────────────────────────────
slide_title(doc, "Frontend — ECMWF GraphCast UI", "React + shadcn/ui + Recharts")

section_heading(doc, "Layout (mirrors charts.ecmwf.int)")
add_code_block(doc,
"""┌─────────────────────────────────────────────────────────────────┐
│  NAVBAR  NeuralWeather | EXPERIMENTAL · Graph-LSTM · Punjab     │
├──────────────────┬──────────────────────────────────────────────┤
│  SELECT          │  Temperature · Ropar, Punjab                 │
│  DIMENSIONS      │                                              │
│  ─────────       │  [Full-width AreaChart]                      │
│  Location ▼      │   ↑ glowing colour per variable              │
│  Variable ▼      │   Orange ReferenceLine = current step        │
│                  │   Stat pills: Current / Min / Max            │
│  Live Conditions ├──────────────────────────────────────────────┤
│  Temp · Hum ·    │  ⏮ ⏪  ▶/⏸  ⏩ ⏭   ───────●──── T+12h  1× ↺  │
│  Wind · Press    ├──────────────────────────────────────────────┤
│  Model Info      │  FOOTER: Model · Data · License              │
└──────────────────┴──────────────────────────────────────────────┘"""
)

section_heading(doc, "Key Components")
two_col_table(doc,
    ['Component', 'File', 'Role'],
    [
        ['WeatherDashboard', 'WeatherDashboard.jsx', 'State orchestrator — data fetching, playback timer'],
        ['ControlPanel',     'ControlPanel.jsx',     'Sidebar — city/variable select, live conditions, model metadata'],
        ['ForecastChart',    'ForecastCharts.jsx',   'Recharts AreaChart — moving ReferenceLine needle'],
        ['PlaybackBar',      'PlaybackBar.jsx',      'Bottom media controls — shadcn Slider + Select'],
    ]
)

section_heading(doc, "shadcn Components Used")
body_bullet(doc, "Slider (orange scrubber) · Select (city/variable) · Tooltip (button hints) · Separator (sidebar dividers) · Badge (model status)")

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# SLIDE 7 — MLOps Pipeline
# ─────────────────────────────────────────────────────────────────────────────
slide_title(doc, "MLOps Pipeline", "Automated Drift Detection & Fine-tuning")

section_heading(doc, "GitHub Actions — .github/workflows/mlops.yml")
add_code_block(doc,
"""Cron: 0 0 1,15 * *  (1st & 15th of each month at 00:00 UTC)
Runner: self-hosted  (your local i5 machine)

Step 1 ── python -m backend.monitor --threshold 2.0 --days 7
           │  exit 0 = no drift (STOP here — model is healthy)
           │  exit 1 = DRIFT_DETECTED
           │  exit 2 = insufficient log data (STOP)
           ▼
Step 2 ── Fetch last 30 days from Open-Meteo → backend/data/recent_30d.pkl
           ▼
Step 3 ── python -m backend.training.train_model --finetune --epochs 5
           ▼
Step 4 ── dvc add weights → dvc push → git commit .dvc files → git push"""
)

section_heading(doc, "Drift Detection Logic (backend/monitor.py)")
for b in [
    "Every /forecast API call appends predictions to backend/data/forecast_log.jsonl (via forecast_logger.py).",
    "Monitor fetches ACTUAL observations from Open-Meteo for the same past-7-day window.",
    "Computes per-city MAE for temperature_2m across all 7 nodes.",
    "If overall MAE > 2.0°C → print DRIFT_DETECTED → exit 1 → triggers retraining.",
    "Uses the same CITIES and DYNAMIC_FEATURES from shared_config — no skew possible.",
]:
    body_bullet(doc, b)

section_heading(doc, "DVC Model Versioning")
for b in [
    "Model weights (.pt, .pkl) are tracked by DVC, not committed to git directly.",
    ".dvc pointer files are committed — each retraining run creates a new version.",
    "Rollback: git checkout <commit> backend/models/stgnn.pt.dvc → dvc pull",
]:
    body_bullet(doc, b)

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# SLIDE 8 — API Reference
# ─────────────────────────────────────────────────────────────────────────────
slide_title(doc, "API Reference", "FastAPI · uvicorn · localhost:8000")

two_col_table(doc,
    ['Endpoint', 'Method', 'Response', 'Error Codes'],
    [
        ['/health',   'GET', '{ status, ml_model_loaded }',                             '—'],
        ['/current',  'GET', '{ nodes: { City: { temp, hum, wind, … } } }',             '502 upstream fail'],
        ['/forecast', 'GET', '{ metadata, current, forecast: { City: [48 steps] } }',   '502 data · 500 model'],
    ]
)

section_heading(doc, "Error Handling Strategy")
for b in [
    "502 Bad Gateway: Open-Meteo API unreachable or returns malformed data.",
    "500 Internal Server Error: STGNN inference failed (full traceback in uvicorn logs).",
    "Graceful fallback: if model not ready, /forecast returns API-only current conditions (no ML).",
    "Validation guard: validate_city_data() checks all 7 nodes have ≥ 48 rows before inference.",
]:
    body_bullet(doc, b)

section_heading(doc, "5-Point Inference Diagnostics (per request, in uvicorn logs)")
two_col_table(doc,
    ['Check', 'What it detects'],
    [
        ['Input shape log',          'Wrong tensor dimensions (data pipeline bug)'],
        ['Input std per feature',    'std < 0.005 → input data is static/flat (scaler bug)'],
        ['Adjacency edge count',     '0 edges → graph is disconnected (GCN disabled)'],
        ['Raw output std over 48h', 'std < 0.001 → model predicts constants (weight bug)'],
        ['Per-city temp range',      'spread < 0.5°C → flat forecast after inverse-transform'],
    ]
)

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# SLIDE 9 — Known Issues & Debugging
# ─────────────────────────────────────────────────────────────────────────────
slide_title(doc, "Known Issues & Debugging", "Lessons learned from the build")

two_col_table(doc,
    ['Issue', 'Root Cause', 'Fix Applied'],
    [
        ['Flat/linear forecast',     'Residuals near-zero → pred = last_obs + 0×48',       'Tanh decoder + ×0.05 scale + clamp(0,1)'],
        ['NumPy 2.x C-ABI crash',   'PyTorch 2.1.0 compiled against NumPy 1.x ABI',       'Pin numpy<2.0.0 in requirements.txt + Dockerfile'],
        ['CORS Network Error',       'FastAPI origins list missing Vite port 5173',          'allow_origins=["*"] for local dev'],
        ['Training-Serving Skew',   'DYNAMIC_FEATURES defined in 5 separate files',         'backend/shared_config.py single source of truth'],
        ['Flat precipitation input', 'Dry season = all zeros → std≈0 → no gradient',        'Diagnostic now flags ⚠ FLAT in logs'],
        ['PyTorch cross-version .pt','2.7.1 saved, 2.1.0 loaded → weights_only clash',      'Use weights_only=False for cross-version loads'],
        ['Tailwind white/8 class',   'Tailwind v3 opacity scale has no /8 step',            'Globally replaced all white/8 with white/10 in JSX'],
    ]
)

section_heading(doc, "Cross-Machine Training Workflow")
for b in [
    "Train on GPU machine (any PyTorch version with CUDA) → copy 2 files to inference machine.",
    "Only stgnn.pt (weights) + stgnn_scaler.pkl (scalers) need to be transferred.",
    "map_location='cpu' in torch.load() handles GPU→CPU automatically.",
    "pkl files are Python-version-tolerant between 3.11 and 3.13 for sklearn scalers.",
]:
    body_bullet(doc, b)

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# SLIDE 10 — Results & Demo
# ─────────────────────────────────────────────────────────────────────────────
slide_title(doc, "Results & Live Demo", "What the system produces")

section_heading(doc, "System Outputs")
two_col_table(doc,
    ['Output', 'Description'],
    [
        ['48-h Temperature Forecast',  'Per-city animated chart from current conditions'],
        ['48-h Humidity Forecast',     'Relative humidity trajectory for 7 nodes'],
        ['48-h Wind Speed Forecast',   'Surface wind speed prediction (km/h)'],
        ['48-h Pressure Forecast',     'Mean sea-level pressure (hPa)'],
        ['Live Conditions Sidebar',    'Real-time observed values from Open-Meteo'],
        ['T+Nh Playback',              'Step through each forecast hour like a weather map'],
    ]
)

section_heading(doc, "Performance (CPU — Intel Core i5)")
two_col_table(doc,
    ['Operation', 'Time'],
    [
        ['Model load at startup',          '~2–4 seconds (PyG DenseGCNConv + LSTMCell)'],
        ['API data fetch (7 cities)',       '~1–2 seconds (7 async httpx calls)'],
        ['STGNN inference (48h × 7 nodes)','~0.1–0.3 seconds'],
        ['Total /forecast latency',        '~2–4 seconds end-to-end'],
        ['Cold-start training (1 year)',    '~30–60 minutes on CPU'],
        ['Fine-tuning (30 days, 5 epochs)', '~5–15 minutes on CPU'],
    ]
)

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# SLIDE 11 — Project Structure
# ─────────────────────────────────────────────────────────────────────────────
slide_title(doc, "Project Structure", "Repository layout")

add_code_block(doc,
"""Weather-Forecast/
├── backend/
│   ├── app/
│   │   ├── main.py                  ← FastAPI app + CORS
│   │   ├── routes/api.py            ← /health /current /forecast
│   │   └── services/
│   │       ├── ml_service.py        ← STGNN inference singleton
│   │       ├── weather_api.py       ← Open-Meteo async fetcher
│   │       └── forecast_logger.py   ← JSONL prediction logger
│   ├── models/
│   │   ├── stgnn.py                 ← WeatherSTGNN architecture
│   │   ├── stgnn.pt                 ← trained weights (DVC)
│   │   └── stgnn_scaler.pkl         ← fitted MinMaxScalers (DVC)
│   ├── training/
│   │   ├── train_model.py           ← cold-start + fine-tune CLI
│   │   ├── data_collector.py        ← Open-Meteo historical fetch
│   │   └── preprocess.py            ← sliding window + scaling
│   ├── data/
│   │   └── forecast_log.jsonl       ← runtime prediction log (gitignored)
│   ├── shared_config.py             ← single source of truth constants
│   ├── monitor.py                   ← drift detection script
│   └── requirements.txt
├── frontend/
│   ├── src/
│   │   ├── components/
│   │   │   ├── WeatherDashboard.jsx
│   │   │   ├── ControlPanel.jsx
│   │   │   ├── ForecastCharts.jsx
│   │   │   └── PlaybackBar.jsx
│   │   ├── services/api.js
│   │   └── index.css                ← ECMWF dark theme tokens
│   └── vite.config.js               ← proxy /forecast → :8000
├── .github/workflows/
│   ├── retrain.yml                  ← daily retraining (GitHub cloud)
│   └── mlops.yml                    ← drift + fine-tune (self-hosted)
├── docker-compose.yml
└── .gitignore""")

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# SLIDE 12 — Future Work
# ─────────────────────────────────────────────────────────────────────────────
slide_title(doc, "Future Work & Extensions", "What comes next")

two_col_table(doc,
    ['Extension', 'Technical Approach', 'Difficulty'],
    [
        ['True GraphCast backbone',          'Replace LSTM with Transformer processor; use icosahedral mesh', 'High'],
        ['Pressure-level 3D input',          'Add 13 ERA5 pressure levels as additional node features',   'Medium'],
        ['Quantile uncertainty bands',        'Train model to output 10th/50th/90th percentile predictions', 'Medium'],
        ['Longer horizon (10 days)',          'Increase FUTURE_STEPS to 240; retrain with larger dataset',  'Low'],
        ['Additional cities / regions',       'Add CITIES to shared_config; rebuild graph edges',          'Low'],
        ['Real Leaflet/Mapbox map overlay',   'Replace AreaChart with geographic weather map render',       'Medium'],
        ['Mobile-responsive UI',              'shadcn Sheet for sidebar on small screens',                  'Low'],
        ['Alert notifications',              'Email/Telegram when drift detected, temperature extreme hit', 'Low'],
    ]
)

section_heading(doc, "Acknowledgements")
for b in [
    "ECMWF — for open-sourcing GraphCast and inspiring this project's UI design",
    "Open-Meteo — for providing free, high-quality hourly weather data",
    "PyTorch Geometric team — for DenseGCNConv and the GNN ecosystem",
    "shadcn/ui — for the beautiful, accessible React component primitives",
]:
    body_bullet(doc, b)

# ─────────────────────────────────────────────────────────────────────────────
# SAVE
# ─────────────────────────────────────────────────────────────────────────────
out = "d:/disk d/Weather-Forecast/NeuralWeather_Project_Presentation.docx"
doc.save(out)
print(f"[OK] Saved -> {out}")
